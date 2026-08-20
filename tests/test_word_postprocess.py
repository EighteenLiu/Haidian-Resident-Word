from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.word_postprocess import HorizontalMergeRule, MergeRule, ParagraphReplacementRule, ParagraphStyleRule, apply_word_postprocess, merge_vertical_same_cells


def test_merge_vertical_same_cells(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "镇"
    table.cell(0, 1).text = "村"
    table.cell(1, 0).text = "上庄镇"
    table.cell(1, 1).text = "八家村"
    table.cell(2, 0).text = "上庄镇"
    table.cell(2, 1).text = "皂甲屯村"
    table.cell(3, 0).text = "苏家坨镇"
    table.cell(3, 1).text = "柳林村"
    doc.save(source)
    merge_vertical_same_cells(source, output, [MergeRule(table_index=0, column_index=0, start_row=1, key_columns=(0,))])
    assert output.exists()
    checked = Document(output)
    assert checked.tables[0].cell(1, 0).text == "上庄镇"
    assert checked.tables[0].cell(1, 0)._tc.tcPr.vAlign.get(qn("w:val")) == "center"


def test_apply_word_postprocess_merges_total_row(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "类型"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "农村生活垃圾治理"
    table.cell(2, 0).text = "总计"
    table.cell(2, 1).text = "总计"
    table.cell(2, 2).text = "663"
    table.cell(2, 3).text = "100%"
    table.cell(2, 0).paragraphs[0].runs[0].bold = True
    table.cell(2, 0).paragraphs[0].runs[0].font.name = "仿宋"
    doc.save(source)

    apply_word_postprocess(
        source,
        output,
        [],
        [
            HorizontalMergeRule(
                table_index=0,
                row_index=-1,
                start_column=0,
                end_column=1,
                expected_texts=("总计", "总计"),
                text="总计",
                bold_row=True,
            )
        ],
    )

    checked = Document(output)
    total_cell = checked.tables[0].cell(2, 0)
    assert total_cell.text == "总计"
    assert total_cell.paragraphs[0].runs[0].bold is True
    assert checked.tables[0].cell(2, 2).paragraphs[0].runs[0].bold is True
    assert total_cell._tc.tcPr.vAlign.get(qn("w:val")) == "center"
    grid_span = total_cell._tc.tcPr.find(qn("w:gridSpan"))
    assert grid_span is not None
    assert grid_span.get(qn("w:val")) == "2"


def test_apply_word_postprocess_merges_three_total_cells(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=5)
    table.cell(1, 0).text = "总计"
    table.cell(1, 1).text = "总计"
    table.cell(1, 2).text = "总计"
    table.cell(1, 3).text = "663"
    table.cell(1, 4).text = "100.00%"
    doc.save(source)

    apply_word_postprocess(
        source,
        output,
        [],
        [
            HorizontalMergeRule(
                table_index=0,
                row_index=-1,
                start_column=0,
                end_column=2,
                expected_texts=("总计", "总计", "总计"),
                text="总计",
                bold_row=True,
            )
        ],
    )

    checked = Document(output)
    total_cell = checked.tables[0].cell(1, 0)
    assert total_cell.text == "总计"
    assert total_cell._tc.tcPr.vAlign.get(qn("w:val")) == "center"
    assert total_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert checked.tables[0].cell(1, 3).paragraphs[0].runs[0].bold is True
    grid_span = total_cell._tc.tcPr.find(qn("w:gridSpan"))
    assert grid_span is not None
    assert grid_span.get(qn("w:val")) == "3"


def test_apply_word_postprocess_replaces_paragraph_text(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("表3  2026年6月农村人居环境检查发现各镇问题数量统计表")
    doc.save(source)

    apply_word_postprocess(
        source,
        output,
        [],
        [],
        [
            ParagraphReplacementRule(
                old="农村人居环境检查发现各镇问题数量统计表",
                new="农村人居环境检查发现各村问题数量统计表",
            )
        ],
    )

    checked = Document(output)
    assert checked.paragraphs[0].text == "表3  2026年6月农村人居环境检查发现各村问题数量统计表"


def test_apply_word_postprocess_styles_table_three_title(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("表3  2026年7月农村人居环境检查发现各村问题数量统计表")
    doc.save(source)

    apply_word_postprocess(
        source,
        output,
        [],
        [],
        [],
        [
            ParagraphStyleRule(
                text="农村人居环境检查发现各村问题数量统计表",
                font_name="方正小标宋简体",
                size_pt=14,
                bold=False,
                align="center",
            )
        ],
    )

    checked = Document(output)
    paragraph = checked.paragraphs[0]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.runs[0].bold is False
    assert paragraph.runs[0].font.size.pt == 14
    assert paragraph.runs[0].font.name == "方正小标宋简体"
