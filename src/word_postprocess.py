from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class MergeRule:
    table_index: int
    column_index: int
    start_row: int = 1
    key_columns: tuple[int, ...] = (0,)
    merge_empty: bool = False
    table_name: str = ""


@dataclass
class HorizontalMergeRule:
    table_index: int
    row_index: int
    start_column: int
    end_column: int
    expected_texts: tuple[str, ...] = ()
    text: str = ""
    bold_row: bool = False
    table_name: str = ""


@dataclass
class ParagraphReplacementRule:
    old: str
    new: str


def _cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _merge_key(row, rule: MergeRule) -> str:
    return "|".join(_cell_text(row.cells[idx]) for idx in rule.key_columns if idx < len(row.cells))


def _center_cell(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "center")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_cell_text_preserving_style(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    _center_cell(cell)


def _bold_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            run.bold = True


def _bold_row(row) -> None:
    seen: set[int] = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        _bold_cell(cell)


def _center_all_table_cells(doc: Document) -> None:
    for table in doc.tables:
        seen: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                _center_cell(cell)


def merge_same_cells_in_column(table, rule: MergeRule) -> None:
    if len(table.rows) <= rule.start_row:
        return
    group_start = rule.start_row
    last_key = _merge_key(table.rows[group_start], rule)
    for row_idx in range(rule.start_row + 1, len(table.rows) + 1):
        current_key = _merge_key(table.rows[row_idx], rule) if row_idx < len(table.rows) else None
        if current_key != last_key:
            group_end = row_idx - 1
            if group_end > group_start and (last_key or rule.merge_empty):
                top = table.cell(group_start, rule.column_index)
                bottom = table.cell(group_end, rule.column_index)
                merged = top.merge(bottom)
                _set_cell_text_preserving_style(merged, last_key.split("|")[0])
            group_start = row_idx
            last_key = current_key


def merge_cells_in_row(table, rule: HorizontalMergeRule) -> None:
    row_index = rule.row_index if rule.row_index >= 0 else len(table.rows) + rule.row_index
    if row_index < 0 or row_index >= len(table.rows):
        return
    if rule.start_column < 0 or rule.end_column >= len(table.rows[row_index].cells) or rule.end_column <= rule.start_column:
        return
    cells = table.rows[row_index].cells
    if rule.expected_texts:
        actual = tuple(_cell_text(cells[idx]) for idx in range(rule.start_column, rule.end_column + 1))
        if actual != rule.expected_texts:
            return
    merged = table.cell(row_index, rule.start_column).merge(table.cell(row_index, rule.end_column))
    _set_cell_text_preserving_style(merged, rule.text or _cell_text(merged))
    if rule.bold_row:
        _bold_row(table.rows[row_index])


def apply_word_postprocess(
    docx_path: Path,
    output_path: Path,
    vertical_rules: list[MergeRule],
    horizontal_rules: list[HorizontalMergeRule] | None = None,
    paragraph_replacements: list[ParagraphReplacementRule] | None = None,
) -> None:
    doc = Document(docx_path)
    for rule in paragraph_replacements or []:
        for paragraph in doc.paragraphs:
            if rule.old in paragraph.text:
                paragraph.text = paragraph.text.replace(rule.old, rule.new)
    for rule in vertical_rules:
        if rule.table_index >= len(doc.tables):
            continue
        merge_same_cells_in_column(doc.tables[rule.table_index], rule)
    for rule in horizontal_rules or []:
        if rule.table_index >= len(doc.tables):
            continue
        merge_cells_in_row(doc.tables[rule.table_index], rule)
    _center_all_table_cells(doc)
    doc.save(output_path)


def merge_vertical_same_cells(docx_path: Path, output_path: Path, rules: list[MergeRule]) -> None:
    apply_word_postprocess(docx_path, output_path, rules, [], [])


def rules_from_config(config: dict) -> list[MergeRule]:
    items = config.get("word_postprocess", {}).get("merge_rules", [])
    return [
        MergeRule(
            table_index=int(item["table_index"]),
            column_index=int(item["column_index"]),
            start_row=int(item.get("start_row", 1)),
            key_columns=tuple(int(x) for x in item.get("key_columns", [0])),
            merge_empty=bool(item.get("merge_empty", False)),
            table_name=str(item.get("table_name", "")),
        )
        for item in items
    ]


def horizontal_rules_from_config(config: dict) -> list[HorizontalMergeRule]:
    items = config.get("word_postprocess", {}).get("horizontal_merge_rules", [])
    return [
        HorizontalMergeRule(
            table_index=int(item["table_index"]),
            row_index=int(item["row_index"]),
            start_column=int(item["start_column"]),
            end_column=int(item["end_column"]),
            expected_texts=tuple(str(x) for x in item.get("expected_texts", [])),
            text=str(item.get("text", "")),
            bold_row=bool(item.get("bold_row", False)),
            table_name=str(item.get("table_name", "")),
        )
        for item in items
    ]


def paragraph_replacements_from_config(config: dict) -> list[ParagraphReplacementRule]:
    items = config.get("word_postprocess", {}).get("paragraph_replacements", [])
    return [
        ParagraphReplacementRule(
            old=str(item["old"]),
            new=str(item["new"]),
        )
        for item in items
    ]
